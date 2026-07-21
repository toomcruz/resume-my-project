from __future__ import annotations

import argparse
import re
import unicodedata
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)


def normalize(value: str) -> str:
    return "".join(
        char
        for char in unicodedata.normalize("NFD", value.lower())
        if unicodedata.category(char) != "Mn"
    )


def iter_block_items(parent: DocxDocument | _Cell) -> Iterable[Paragraph | Table]:
    parent_element = parent.element.body if isinstance(parent, DocxDocument) else parent._tc
    for child in parent_element.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, parent)
        elif child.tag == qn("w:tbl"):
            yield Table(child, parent)


def iter_tables(document: DocxDocument) -> Iterable[Table]:
    seen: set[int] = set()

    def walk(container: DocxDocument | _Cell):
        for block in iter_block_items(container):
            if not isinstance(block, Table):
                continue
            if id(block._tbl) not in seen:
                seen.add(id(block._tbl))
                yield block
            for row in block.rows:
                for cell in row.cells:
                    yield from walk(cell)

    yield from walk(document)
    for section in document.sections:
        for container in (section.header, section.footer):
            for table in container.tables:
                if id(table._tbl) not in seen:
                    seen.add(id(table._tbl))
                    yield table


def iter_paragraphs(document: DocxDocument) -> Iterable[Paragraph]:
    seen: set[int] = set()

    def walk(container):
        for paragraph in container.paragraphs:
            if id(paragraph._p) not in seen:
                seen.add(id(paragraph._p))
                yield paragraph
        for table in container.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from walk(cell)

    yield from walk(document)
    for section in document.sections:
        yield from walk(section.header)
        yield from walk(section.footer)


def set_cell_margins(cell: _Cell, top=90, start=110, bottom=90, end=110):
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell: _Cell, color="7F7F7F", size=8, style="single"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)
        node.set(qn("w:space"), "0")


def shade_cell(cell: _Cell, fill: str):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_table_width(table: Table, percentage=100):
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(percentage * 50))
    width.set(qn("w:type"), "pct")
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "autofit")
    table.autofit = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def set_page(document: DocxDocument, landscape=False, margins=(1.2, 1.2, 1.2, 1.2)):
    for section in document.sections:
        if landscape:
            section.orientation = WD_ORIENT.LANDSCAPE
            section.page_width = A4_HEIGHT
            section.page_height = A4_WIDTH
        else:
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = A4_WIDTH
            section.page_height = A4_HEIGHT
        top, right, bottom, left = margins
        section.top_margin = Cm(top)
        section.right_margin = Cm(right)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)


def remove_fixed_heights(table: Table):
    for row in table.rows:
        row.height = None
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        properties = row._tr.get_or_add_trPr()
        for tag in ("w:trHeight", "w:cantSplit"):
            node = properties.find(qn(tag))
            if node is not None:
                properties.remove(node)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            no_wrap = cell._tc.get_or_add_tcPr().find(qn("w:noWrap"))
            if no_wrap is not None:
                cell._tc.get_or_add_tcPr().remove(no_wrap)


def normalize_typography(document: DocxDocument, minimum=9.0, footer_minimum=9.0):
    for paragraph in iter_paragraphs(document):
        paragraph.paragraph_format.space_before = Pt(0)
        if paragraph.paragraph_format.space_after is None or paragraph.paragraph_format.space_after.pt > 6:
            paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        if paragraph.paragraph_format.left_indent is not None and paragraph.paragraph_format.left_indent.pt < 0:
            paragraph.paragraph_format.left_indent = Pt(0)
        if paragraph.paragraph_format.right_indent is not None and paragraph.paragraph_format.right_indent.pt < 0:
            paragraph.paragraph_format.right_indent = Pt(0)
        for run in paragraph.runs:
            if run.font.size is not None and run.font.size.pt < minimum:
                run.font.size = Pt(minimum)
            if run.font.name is None:
                run.font.name = "Arial"
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Arial")
    for section in document.sections:
        for paragraph in section.footer.paragraphs:
            for run in paragraph.runs:
                if run.font.size is None or run.font.size.pt < footer_minimum:
                    run.font.size = Pt(footer_minimum)


def scale_inline_images(document: DocxDocument, maximum_cm: float):
    maximum = int(Cm(maximum_cm))
    for inline in document._element.xpath(".//wp:inline"):
        extent = inline.find(qn("wp:extent"))
        if extent is None:
            continue
        width = int(extent.get("cx", "0"))
        height = int(extent.get("cy", "0"))
        if width <= maximum or width <= 0:
            continue
        ratio = maximum / width
        extent.set("cx", str(maximum))
        extent.set("cy", str(int(height * ratio)))


def common_fix(document: DocxDocument, landscape=False, minimum=9.0):
    set_page(document, landscape=landscape)
    normalize_typography(document, minimum=minimum)
    for table in iter_tables(document):
        set_table_width(table)
        remove_fixed_heights(table)
    scale_inline_images(document, 26 if landscape else 18)


def replace_paragraph_text(document: DocxDocument, replacements: list[tuple[re.Pattern, str]]):
    for paragraph in iter_paragraphs(document):
        if not paragraph.runs:
            continue
        original = "".join(run.text for run in paragraph.runs)
        updated = original
        for pattern, replacement in replacements:
            updated = pattern.sub(replacement, updated)
        if updated == original:
            continue
        paragraph.runs[0].text = updated
        for run in paragraph.runs[1:]:
            run.text = ""


def insert_banner(document: DocxDocument, text: str):
    table = document.add_table(rows=1, cols=1)
    set_table_width(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "D9EAF7")
    set_cell_border(cell, color="2F75B5", size=10)
    set_cell_margins(cell, 110, 150, 110, 150)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    body = document._element.body
    body.remove(table._tbl)
    body.insert(0, table._tbl)


def extract_placeholders(document: DocxDocument) -> list[str]:
    text = "\n".join(paragraph.text for paragraph in iter_paragraphs(document))
    result: list[str] = []
    for value in re.findall(r"\{\{?\s*([A-Za-zÀ-ÿ0-9_]+)\s*\}?\}", text):
        if value not in result:
            result.append(value)
    return result


def choose_placeholder(values: list[str], hints: tuple[str, ...], fallback: str) -> str:
    for hint in hints:
        for value in values:
            if hint in normalize(value):
                return "{{" + value + "}}"
    return "{{" + fallback + "}}"


def rebuild_condolences(source: Path, target: Path):
    original = Document(source)
    fields = extract_placeholders(original)
    deceased = choose_placeholder(fields, ("nomefal", "falecido", "nome"), "nomeFal")
    date = choose_placeholder(fields, ("datasep", "data"), "dataSep")
    hour = choose_placeholder(fields, ("horasep", "hora"), "horaSep")

    document = Document()
    set_page(document, margins=(1.5, 1.5, 1.5, 1.5))
    frame = document.add_table(rows=1, cols=1)
    set_table_width(frame)
    cell = frame.cell(0, 0)
    set_cell_border(cell, color="707070", size=12, style="dashed")
    set_cell_margins(cell, 350, 350, 350, 350)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = paragraph.add_run("MESA DE CONDOLÊNCIAS")
    title.bold = True
    title.font.size = Pt(20)
    paragraph.paragraph_format.space_after = Pt(16)

    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name = paragraph.add_run(deceased)
    name.bold = True
    name.font.size = Pt(25)
    paragraph.paragraph_format.space_after = Pt(16)

    information = cell.add_table(rows=2, cols=2)
    set_table_width(information, 80)
    for index, (label, value) in enumerate((("DATA", date), ("HORÁRIO", hour))):
        label_cell, value_cell = information.rows[index].cells
        label_cell.text = label
        value_cell.text = value
        shade_cell(label_cell, "EAF2F8")
        for current in (label_cell, value_cell):
            set_cell_border(current)
            set_cell_margins(current, 160, 160, 160, 160)
            current.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in label_cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(12)
        for run in value_cell.paragraphs[0].runs:
            run.font.size = Pt(14)

    note = cell.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(18)
    note_run = note.add_run("Registre sua mensagem e assinatura.")
    note_run.italic = True
    note_run.font.size = Pt(11)
    for _ in range(8):
        line = cell.add_paragraph("________________________________________________________________________")
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line.paragraph_format.space_after = Pt(4)
        for run in line.runs:
            run.font.size = Pt(10)
    document.save(target)


def rebuild_room_identification(source: Path, target: Path):
    original = Document(source)
    fields = extract_placeholders(original)
    deceased = choose_placeholder(fields, ("nomefal", "falecido", "nome"), "nomeFal")
    room = choose_placeholder(fields, ("sala",), "salaVelorio")
    date = choose_placeholder(fields, ("datasep", "data"), "dataSep")
    hour = choose_placeholder(fields, ("horasep", "hora"), "horaSep")

    document = Document()
    set_page(document, landscape=True, margins=(1.2, 1.6, 1.2, 1.6))
    frame = document.add_table(rows=1, cols=1)
    set_table_width(frame)
    cell = frame.cell(0, 0)
    set_cell_border(cell, color="1F4E78", size=18)
    set_cell_margins(cell, 450, 400, 450, 400)

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("IDENTIFICAÇÃO DA SALA DE VELÓRIO")
    run.bold = True
    run.font.size = Pt(19)
    paragraph.paragraph_format.space_after = Pt(18)

    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(deceased)
    run.bold = True
    run.font.size = Pt(29)
    paragraph.paragraph_format.space_after = Pt(20)

    paragraph = cell.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(f"SALA  {room}")
    run.bold = True
    run.font.size = Pt(23)
    paragraph.paragraph_format.space_after = Pt(18)

    information = cell.add_table(rows=1, cols=2)
    set_table_width(information, 75)
    for index, (label, value) in enumerate((("DATA", date), ("HORÁRIO", hour))):
        current = information.cell(0, index)
        set_cell_border(current)
        set_cell_margins(current, 200, 200, 200, 200)
        current.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = current.paragraphs[0].add_run(f"{label}:  {value}")
        run.bold = True
        run.font.size = Pt(16)
    document.save(target)


def fix_memorandum(document: DocxDocument):
    common_fix(document)
    insert_banner(document, "MEMORANDO DE TRANSLADO · FORMATO A4")


def fix_order(document: DocxDocument, process: str, place: str):
    common_fix(document, minimum=8.5)
    insert_banner(document, f"ORDEM DE {process.upper()} · {place.upper()}")
    if place == "Jazigo":
        replace_paragraph_text(
            document,
            [
                (re.compile(r"QUADRA\s+GERAL", re.I), "JAZIGO"),
                (re.compile(r"Quadra\s+Geral", re.I), "Jazigo"),
            ],
        )
    else:
        replace_paragraph_text(document, [(re.compile(r"\bJAZIGO\b", re.I), "QUADRA GERAL")])
    for table in iter_tables(document):
        for row in table.rows:
            for cell in row.cells:
                compact = normalize(cell.text).replace(" ", "")
                if any(term in compact for term in ("quadrarua", "quadra/rua", "terreno", "jazigo")):
                    cell.width = Cm(3.4)
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_after = Pt(0)
                        for run in paragraph.runs:
                            run.font.size = Pt(8.5)


def fix_ossuary(document: DocxDocument, mode: str):
    common_fix(document)
    replace_paragraph_text(
        document,
        [(re.compile(r"(?<!\{)\{([A-Za-zÀ-ÿ0-9_]+)\}(?!\})"), r"{{\1}}")],
    )
    acquisition = mode == "Aquisição"
    banner = (
        "☒ AQUISIÇÃO / 1º ALUGUEL     ☐ RENOVAÇÃO"
        if acquisition
        else "☐ AQUISIÇÃO / 1º ALUGUEL     ☒ RENOVAÇÃO"
    )
    insert_banner(document, banner)

    text = "\n".join(paragraph.text for paragraph in iter_paragraphs(document))
    if "inscrGS" not in text:
        paragraph_element = OxmlElement("w:p")
        document._element.body.insert(1, paragraph_element)
        paragraph = Paragraph(paragraph_element, document)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("INSCRIÇÃO GSCEMI: {{inscrGS}}")
        run.bold = True
        run.font.size = Pt(10)

    for table in iter_tables(document):
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell, 105, 155, 105, 155)
                cell_text = normalize(cell.text)
                if "data" in cell_text or "vencimento" in cell_text:
                    cell.width = Cm(4)
                    properties = cell._tc.get_or_add_tcPr()
                    if properties.find(qn("w:noWrap")) is None:
                        properties.append(OxmlElement("w:noWrap"))


def fix_term(document: DocxDocument, process: str):
    common_fix(document)
    checked = (
        "☒ SEPULTAMENTO     ☐ EXUMAÇÃO"
        if process == "Sepultamento"
        else "☐ SEPULTAMENTO     ☒ EXUMAÇÃO"
    )
    insert_banner(document, f"TERMO DE COMPROMISSO E RESPONSABILIDADE · {checked}")
    replace_paragraph_text(
        document,
        [
            (
                re.compile(
                    r"[\(\[]\s*[ xX]?\s*[\)\]]\s*SEPULTAMENTO\s+[\(\[]\s*[ xX]?\s*[\)\]]\s*EXUMAÇÃO",
                    re.I,
                ),
                checked,
            ),
            (re.compile(r"SEPULTAMENTO\s*/\s*EXUMAÇÃO", re.I), process.upper()),
        ],
    )


def classify(path: Path) -> tuple[str, str | None]:
    name = normalize(path.stem)
    if "condol" in name:
        return "condolences", None
    if "identificacao" in name and "velorio" in name:
        return "identification", None
    if "memorando" in name or "translado" in name:
        return "memorandum", None
    if "atualizacao" in name and "cadastral" in name:
        return "registration", None
    if "semi" in name and "intacto" in name:
        return "semi-intact", None
    if "ordem" in name and "exum" in name:
        return "exhumation-order", "Jazigo" if "jazigo" in name else "Quadra Geral"
    if "ordem" in name and "sepult" in name:
        return "burial-order", "Jazigo" if "jazigo" in name else "Quadra Geral"
    if "ossuario" in name:
        return "ossuary", "Renovação" if "renov" in name else "Aquisição"
    if "termo" in name and "respons" in name:
        return "term", "Exumação" if "exum" in name else "Sepultamento"
    return "generic", None


def fix_file(source: Path, target: Path):
    kind, variant = classify(source)
    if kind == "condolences":
        rebuild_condolences(source, target)
        return
    if kind == "identification":
        rebuild_room_identification(source, target)
        return

    document = Document(source)
    if kind == "memorandum":
        fix_memorandum(document)
    elif kind == "registration":
        common_fix(document)
    elif kind == "semi-intact":
        common_fix(document)
        scale_inline_images(document, 17.5)
    elif kind == "exhumation-order":
        fix_order(document, "Exumação", variant or "Quadra Geral")
    elif kind == "burial-order":
        fix_order(document, "Sepultamento", variant or "Quadra Geral")
    elif kind == "ossuary":
        fix_ossuary(document, variant or "Aquisição")
    elif kind == "term":
        fix_term(document, variant or "Sepultamento")
    else:
        common_fix(document)
    document.save(target)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source", type=Path)
    parser.add_argument("target", type=Path)
    arguments = parser.parse_args()
    arguments.target.mkdir(parents=True, exist_ok=True)

    files = sorted(arguments.source.glob("*.docx"))
    if not files:
        raise SystemExit(f"Nenhum DOCX encontrado em {arguments.source}")

    for source in files:
        target = arguments.target / source.name
        fix_file(source, target)
        print(f"CORRIGIDO {source.name}")

    print(f"TOTAL_CORRIGIDO={len(files)}")


if __name__ == "__main__":
    main()
